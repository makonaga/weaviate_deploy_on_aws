import json
import boto3
import os
import logging
from io import BytesIO
from PyPDF2 import PdfReader
import docx
import requests
import uuid
from urllib.parse import unquote_plus  # ✅ Fix for S3 key decoding

# Config from environment
WEAVIATE_URL = os.environ.get("WEAVIATE_URL")
BEDROCK_REGION = os.environ.get("BEDROCK_REGION")
EMBEDDING_MODEL = os.environ.get("EMBEDDING_MODEL")

logger = logging.getLogger()
logger.setLevel(logging.INFO)

s3 = boto3.client("s3")
bedrock = boto3.client("bedrock-runtime", region_name=BEDROCK_REGION)

def extract_text_from_pdf(file_stream):
    logger.info("Extracting text from PDF...")
    reader = PdfReader(file_stream)
    return "".join([page.extract_text() or "" for page in reader.pages])

def extract_text_from_docx(file_stream):
    logger.info("Extracting text from DOCX...")
    document = docx.Document(file_stream)
    text = "\n".join(para.text for para in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    return text

def extract_text_from_txt(file_stream):
    logger.info("Extracting text from TXT...")
    return file_stream.read().decode("utf-8", errors="ignore")

def get_embeddings(text_chunk):
    logger.info("Generating embedding from Titan for chunk...")
    body = {"inputText": text_chunk}
    response = bedrock.invoke_model(
        modelId=EMBEDDING_MODEL,
        body=json.dumps(body),
        contentType="application/json"
    )
    response_body = json.loads(response["body"].read())
    return response_body["embedding"]

def chunk_text(text, chunk_size=500, overlap=100):
    logger.info(f"Chunking text: total length = {len(text)}")
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end])
        start += chunk_size - overlap
    logger.info(f"Created {len(chunks)} chunks")
    return chunks

def extract_context_details_from_key(key):
    """
    Parses keys like:
    context/COOLNN__question_generation__Deployment-Guide/Filename.docx
    """
    try:
        folder = key.split("/")[1]  # "COOLNN__question_generation__Deployment-Guide"
        context_id, context_type, context_name = folder.split("__")
        return context_id, context_type, context_name
    except Exception as e:
        raise ValueError(f"Invalid S3 key format for context extraction: {e}")

def lambda_handler(event, context):
    logger.info("Lambda triggered")
    logger.info("Received event: %s", json.dumps(event))

    try:
        bucket = event["Records"][0]["s3"]["bucket"]["name"]
        encoded_key = event["Records"][0]["s3"]["object"]["key"]
        key = unquote_plus(encoded_key)

        extension = os.path.splitext(key)[1].lower()
        logger.info(f"Bucket: {bucket}, Key: {key}, Extension: {extension}")

        obj = s3.get_object(Bucket=bucket, Key=key)
        file_bytes = obj["Body"].read()
        file_stream = BytesIO(file_bytes)

        logger.info(f"File downloaded: size = {len(file_bytes)} bytes")

        if extension == ".pdf":
            text = extract_text_from_pdf(file_stream)
        elif extension == ".docx":
            text = extract_text_from_docx(file_stream)
        elif extension == ".txt":
            text = extract_text_from_txt(file_stream)
        else:
            logger.error("Unsupported file type")
            return {"statusCode": 400, "body": json.dumps("Unsupported file type")}

        logger.info(f"Extracted text length = {len(text)}")
        chunks = chunk_text(text)

        context_id, context_type, context_name = extract_context_details_from_key(key)
        logger.info(f"Context ID: {context_id}, Type: {context_type}, Name: {context_name}")

        for i, chunk in enumerate(chunks):
            embedding = get_embeddings(chunk)

            payload = {
                "class": "DocumentChunk",
                "id": str(uuid.uuid4()),
                "vector": embedding,
                "properties": {
                    "text": chunk,
                    "context_id": context_id,
                    "context_type": context_type,
                    "context_name": context_name,
                    "metadata": key
                }
            }

            logger.info(f"Pushing chunk {i+1}/{len(chunks)} to Weaviate")
            res = requests.post(
                f"{WEAVIATE_URL}/v1/objects",
                headers={"Content-Type": "application/json"},
                data=json.dumps(payload)
            )

            logger.info(f"Chunk {i+1} Weaviate response: {res.status_code} - {res.text}")

        return {
            "statusCode": 200,
            "body": json.dumps({
                "message": f"{len(chunks)} chunks embedded and uploaded",
                "context_id": context_id,
                "context_type": context_type,
                "context_name": context_name
            })
        }

    except Exception as e:
        logger.error("Exception occurred: %s", str(e), exc_info=True)
        return {
            "statusCode": 500,
            "body": json.dumps("Failed to process document")
        }
